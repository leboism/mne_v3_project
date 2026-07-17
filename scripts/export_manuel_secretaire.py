#!/usr/bin/env python3
"""Exporte le manuel secrétariat en .docx et .pdf (images incluses)."""

from __future__ import annotations

import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
MD_PATH = DOCS / "manuel_secretaire.md"
IMG_ROOT = DOCS


def _parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if not line.startswith("|"):
        return []
    cells = [c.strip() for c in line.strip("|").split("|")]
    return cells


def _is_table_sep(line: str) -> bool:
    return bool(re.match(r"^\|[\s\-:|]+\|\s*$", line.strip()))


def export_docx(path: Path) -> None:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    text = MD_PATH.read_text(encoding="utf-8")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = text.splitlines()
    i = 0
    table_buffer: list[list[str]] | None = None

    def flush_table() -> None:
        nonlocal table_buffer
        if not table_buffer or len(table_buffer) < 1:
            table_buffer = None
            return
        rows = table_buffer
        table_buffer = None
        ncols = max(len(r) for r in rows)
        tbl = doc.add_table(rows=len(rows), cols=ncols)
        tbl.style = "Table Grid"
        for ri, row in enumerate(rows):
            for ci in range(ncols):
                val = row[ci] if ci < len(row) else ""
                tbl.rows[ri].cells[ci].text = val
        doc.add_paragraph()

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            if table_buffer:
                flush_table()
            i += 1
            continue

        if stripped.startswith("|") and _is_table_sep(stripped):
            i += 1
            continue

        if stripped.startswith("|"):
            cells = _parse_table_row(stripped)
            if table_buffer is None:
                table_buffer = []
            table_buffer.append(cells)
            i += 1
            continue

        if table_buffer:
            flush_table()

        img = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if img:
            alt, rel = img.group(1), img.group(2)
            img_path = (IMG_ROOT / rel).resolve()
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if img_path.is_file():
                run = p.add_run()
                run.add_picture(str(img_path), width=Inches(6.2))
            else:
                p.add_run(f"[Image manquante : {rel}]")
            if alt:
                cap = doc.add_paragraph(alt)
                cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cap.style = "Intense Quote"
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=0)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("> "):
            doc.add_paragraph(stripped[2:].strip(), style="Intense Quote")
            i += 1
            continue

        if stripped.startswith("- [ ] "):
            doc.add_paragraph(stripped[6:].strip(), style="List Bullet")
            i += 1
            continue
        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        # Italic caption line *Figure...*
        if stripped.startswith("*") and stripped.endswith("*"):
            p = doc.add_paragraph(stripped.strip("*"))
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            i += 1
            continue

        # Inline **bold** simplified
        plain = re.sub(r"\*\*([^*]+)\*\*", r"\1", stripped)
        plain = re.sub(r"`([^`]+)`", r"\1", plain)
        doc.add_paragraph(plain)
        i += 1

    if table_buffer:
        flush_table()

    doc.save(str(path))
    print("DOCX", path)


def export_pdf(path: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        HRFlowable,
        Image,
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    text = MD_PATH.read_text(encoding="utf-8")
    styles = getSampleStyleSheet()
    h0 = ParagraphStyle("H0", parent=styles["Heading1"], fontSize=18, spaceAfter=12)
    h1 = ParagraphStyle("H1", parent=styles["Heading2"], fontSize=14, spaceBefore=14, spaceAfter=8)
    h2 = ParagraphStyle("H2", parent=styles["Heading3"], fontSize=12, spaceBefore=10, spaceAfter=6)
    body = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10, leading=14, alignment=TA_JUSTIFY)
    quote = ParagraphStyle("Quote", parent=body, textColor=colors.HexColor("#444444"), leftIndent=12)
    caption = ParagraphStyle("Cap", parent=body, fontSize=9, alignment=TA_CENTER, textColor=colors.HexColor("#555555"))
    bullet = ParagraphStyle("Bullet", parent=body, leftIndent=18, bulletIndent=8)

    story: list = []
    lines = text.splitlines()
    i = 0
    table_buffer: list[list[str]] | None = None

    def esc(s: str) -> str:
        return (
            s.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

    def flush_table() -> None:
        nonlocal table_buffer
        if not table_buffer:
            return
        rows = table_buffer
        table_buffer = None
        data = [[Paragraph(esc(c), body) for c in row] for row in rows]
        if not data:
            return
        ncols = len(data[0])
        col_w = (16.0 * cm) / ncols
        t = Table(data, colWidths=[col_w] * ncols)
        t.setStyle(
            TableStyle(
                [
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f0f0")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        story.append(t)
        story.append(Spacer(1, 0.3 * cm))

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            if table_buffer:
                flush_table()
            i += 1
            continue

        if stripped.startswith("|") and _is_table_sep(stripped):
            i += 1
            continue
        if stripped.startswith("|"):
            if table_buffer is None:
                table_buffer = []
            table_buffer.append(_parse_table_row(stripped))
            i += 1
            continue
        if table_buffer:
            flush_table()

        img = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if img:
            alt, rel = img.group(1), img.group(2)
            img_path = IMG_ROOT / rel
            if img_path.is_file():
                im = Image(str(img_path))
                max_w = 16 * cm
                iw, ih = im.imageWidth, im.imageHeight
                if iw > max_w:
                    ratio = max_w / iw
                    im.drawWidth = max_w
                    im.drawHeight = ih * ratio
                story.append(Spacer(1, 0.2 * cm))
                story.append(im)
            if alt:
                story.append(Paragraph(esc(alt), caption))
            story.append(Spacer(1, 0.25 * cm))
            i += 1
            continue

        if stripped.startswith("# "):
            story.append(Spacer(1, 0.2 * cm))
            story.append(Paragraph(esc(stripped[2:]), h0))
            i += 1
            continue
        if stripped.startswith("## "):
            story.append(Paragraph(esc(stripped[3:]), h1))
            i += 1
            continue
        if stripped.startswith("### "):
            story.append(Paragraph(esc(stripped[4:]), h2))
            i += 1
            continue
        if stripped.startswith("> "):
            story.append(Paragraph(esc(stripped[2:]), quote))
            i += 1
            continue
        if stripped.startswith("- "):
            story.append(Paragraph(f"• {esc(stripped[2:])}", bullet))
            i += 1
            continue
        if stripped == "---":
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
            i += 1
            continue
        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            story.append(Paragraph(esc(stripped.strip("*")), caption))
            i += 1
            continue

        plain = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", esc(stripped))
        plain = re.sub(r"`([^`]+)`", r"<font name='Courier'>\1</font>", plain)
        story.append(Paragraph(plain, body))
        i += 1

    if table_buffer:
        flush_table()

    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title="Manuel secrétariat MNE Grade Manager V3",
    )
    doc.build(story)
    print("PDF", path)


def main() -> int:
    if not MD_PATH.is_file():
        print(f"Manuel introuvable : {MD_PATH}", file=sys.stderr)
        return 1

    docx_path = DOCS / "manuel_secretaire.docx"
    pdf_path = DOCS / "manuel_secretaire.pdf"

    try:
        export_docx(docx_path)
    except ImportError:
        print("Installation de python-docx…", file=sys.stderr)
        import subprocess

        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", "python-docx", "-q"],
        )
        export_docx(docx_path)

    export_pdf(pdf_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
